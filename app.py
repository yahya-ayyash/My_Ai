from flask import Flask, render_template, request, session, redirect, url_for, flash
# Importing the tool to translate text (Google Translate)
from deep_translator import GoogleTranslator
import sqlite3
import threading
import webbrowser
import time
# Tools for password security (hashing)
from werkzeug.security import generate_password_hash, check_password_hash
# Tools for making "decorators" (special functions that wrap other functions)
from functools import wraps
import os
# Importing AI providers
import g4f
from pytgpt.phind import PHIND
import sys
import concurrent.futures
import re
import PyPDF2
import docx
from werkzeug.utils import secure_filename
from fpdf import FPDF
import io
from flask import send_file
import requests
from urllib.parse import quote
from groq import Groq

# Initialize Groq client
GROQ_API_KEY = "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"
client = Groq(api_key=GROQ_API_KEY)

# PyInstaller fix for --noconsole: redirect stdout/stderr to a dummy writer if they are None
class DummyWriter:
    def write(self, *args, **kwargs):
        pass
    def flush(self):
        pass

if sys.stdout is None:
    sys.stdout = DummyWriter()
if sys.stderr is None:
    sys.stderr = DummyWriter()

# Initialize the Flask application
# Determine if running as a script or frozen exe
if getattr(sys, 'frozen', False):
    # If frozen, use AppData/Local/CRAB_AI for the database to keep the exe folder clean
    # This prevents the "database.db" file from appearing next to the exe
    app_data = os.environ.get('LOCALAPPDATA', os.path.expanduser('~'))
    BASE_DIR = os.path.join(app_data, 'CRAB_AI')
    
    # Create the directory if it doesn't exist
    if not os.path.exists(BASE_DIR):
        os.makedirs(BASE_DIR)
        
    # And use the temp directory (_MEIPASS) for static/templates
    template_folder = os.path.join(sys._MEIPASS, 'templates')
    static_folder = os.path.join(sys._MEIPASS, 'static')
    app = Flask(__name__, template_folder=template_folder, static_folder=static_folder)
else:
    # If running as script, use the current file's directory
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    app = Flask(__name__)

app.secret_key = 'crab_secret_key_123' # Change this in production

# -------------------------
# DATABASE SETTINGS
# -------------------------
DB_PATH = os.path.join(BASE_DIR, 'database.db')
VERSION = "1.0"
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'csv', 'json', 'py', 'js', 'html', 'css', 'c', 'cpp', 'java'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error extracting PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
    return text

def extract_text_from_file(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext == 'docx':
        return extract_text_from_docx(file_path)
    elif ext == 'csv':
        try:
            import csv
            text = ""
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    text += " | ".join(row) + "\n"
            return text
        except: pass
        
    # Fallback for all text-based files
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except:
            return ""
    except Exception as e:
        print(f"Error reading text file: {e}")
        return ""

def is_small_talk(text):
    """
    Detects simple conversational messages like greetings or thanks.
    """
    if not text: return False
    # Clean the text: lower case, remove punctuation
    clean = text.lower().strip().replace('?', '').replace('!', '').replace('.', '')
    small_talk_phrases = {
        'hi', 'hello', 'hey', 'greetings', 'morning', 'afternoon', 'evening',
        'thanks', 'thank you', 'thx', 'tysm', 'awesome', 'great', 'ok', 'okay',
        'bye', 'goodbye', 'see ya', 'nice', 'cool', 'thank', 'perfect', 'understands', 'understand',
        'got it', 'sure', 'fine', 'no problem', 'youre welcome', 'welcome'
    }
    return clean in small_talk_phrases or len(clean.split()) <= 2

def is_export_request(text):
    """
    Detects if the user is asking to download or export the result.
    """
    if not text: return False
    keywords = ['download', 'export', 'pdf', 'docx', 'word', 'txt', 'save']
    clean = text.lower()
    return any(kw in clean for kw in keywords)

# Inject into jinja templates
app.jinja_env.globals.update(is_export_request=is_export_request)

def init_db():
    """
    Initializes the database by creating necessary tables if they don't exist.
    """
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 1. Users table: Stores username and secure password hash
    cursor.execute('''CREATE TABLE IF NOT EXISTS users 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  username TEXT UNIQUE NOT NULL, 
                  password_hash TEXT NOT NULL)''')
                  
    # 2. Chats table: Stores the list of conversations (threads)
    cursor.execute('''CREATE TABLE IF NOT EXISTS chats 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  user_id INTEGER,
                  title TEXT,
                  timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                  FOREIGN KEY (user_id) REFERENCES users (id))''')
                  
    # 3. History table: Stores individual messages within a chat
    cursor.execute('''CREATE TABLE IF NOT EXISTS history 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  user_id INTEGER,
                  chat_id INTEGER,
                  query TEXT, 
                  result TEXT, 
                  doc_name TEXT,
                  timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                  FOREIGN KEY (user_id) REFERENCES users (id),
                  FOREIGN KEY (chat_id) REFERENCES chats (id))''')
                  
    # 4. Uploads context table: Stores document text temporarily for AI processing
    cursor.execute('''CREATE TABLE IF NOT EXISTS uploads_context 
                 (user_id INTEGER PRIMARY KEY, 
                  filename TEXT, 
                  content TEXT,
                  timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                  FOREIGN KEY (user_id) REFERENCES users (id))''')
                  
    conn.commit() # Save changes
    conn.close()  # Close connection

def get_db_connection():
    """
    Helper function to get a connection to the database.
    """
    conn = sqlite3.connect(DB_PATH, timeout=30)  # Add timeout to prevent locks
    conn.row_factory = sqlite3.Row # Allows accessing columns by name (e.g., row['username'])
    return conn

# -------------------------
# SECURITY & AUTHENTICATION
# -------------------------

# Login Required Decorator
def login_required(view_function):
    """
    This is a 'decorator' that checks if a user is logged in before allowing access to a page.
    """
    @wraps(view_function)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            # Check if it's an background (AJAX) request
            if request.args.get('ajax') == 'true' or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return {"error": "Session expired. Please log in again.", "auth_error": True}, 401
            # If not logged in, redirect to login page
            return redirect(url_for('login'))
        return view_function(*args, **kwargs)
    return decorated_function

# Global Error Handler for AJAX
@app.errorhandler(Exception)
def handle_exception(e):
    """
    Handles errors globally. If it's a background request, return JSON error instead of crashing.
    """
    if request.args.get('ajax') == 'true' or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        print(f"CRITICAL ERROR (AJAX): {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": f"Server error: {str(e)}"}, 500
    return str(e), 500

# -------------------------
# WEBSITE ROUTES (PAGES)
# -------------------------

@app.route("/")
def landing():
    # Shows the landing page (Home)
    return render_template("home.html")

@app.route("/about")
def about():
    # Shows the about page
    return render_template("about.html")

@app.route("/portfolio")
def portfolio():
    # Shows the portfolio page
    return render_template("portfolio.html")

@app.route("/signup", methods=["GET", "POST"])
def signup():
    """
    Handles user registration.
    GET: Show the signup form.
    POST: Process the form data.
    """
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        
        conn = get_db_connection()
        try:
            # Store the user with a HASHED password (never store plain passwords!)
            conn.execute('INSERT INTO users (username, password_hash) VALUES (?, ?)',
                         (username, generate_password_hash(password)))
            conn.commit()
            flash("Account created! Please log in.", "success")
            return redirect(url_for('login'))
        except sqlite3.IntegrityError:
            flash("Username already exists.", "error")
        finally:
            conn.close()
            
    return render_template("signup.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    """
    Handles user login.
    """
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        
        conn = get_db_connection()
        user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
        conn.close()
        
        # Check if user exists AND if the password matches the hash
        if user and check_password_hash(user['password_hash'], password):
            # Save user info in the "session" (browser cookie)
            session['user_id'] = user['id']
            session['username'] = user['username']
            return redirect(url_for('app_interface'))
        else:
            flash("Invalid username or password.", "error")
            
    return render_template("login.html")

@app.route("/logout")
def logout():
    # Clear the session (log out)
    session.clear()
    return redirect(url_for('landing'))

# -------------------------
# APP INTERFACE & AI LOGIC
# -------------------------

@app.route("/app")
@app.route("/app/<int:chat_id>")
@login_required
def app_interface(chat_id=None):
    """
    The main chat interface.
    """
    conn = get_db_connection()
    # Fetch list of previous chats for the sidebar
    chats = conn.execute('SELECT * FROM chats WHERE user_id = ? ORDER BY timestamp DESC', 
                         (session['user_id'],)).fetchall()
    
    active_chat_id = chat_id
    history = []
    
    # If a specific chat is selected, load its message history
    if active_chat_id:
        history = conn.execute('SELECT * FROM history WHERE chat_id = ? ORDER BY timestamp ASC', 
                               (active_chat_id,)).fetchall()
    
    conn.close()
    return render_template("index.html", history=history, chats=chats, active_chat_id=active_chat_id)

@app.route("/new_chat")
@login_required
def new_chat():
    # Start a fresh conversation by removing the chat_id
    return redirect(url_for('app_interface'))

@app.route("/translate", methods=['POST'])
@login_required
def translate():
    """
    Translates text to Tamil using Google Translate.
    """
    data = request.get_json()
    text = data.get("text", "")
    if not text:
        return {"error": "No text provided"}, 400
    
    try:
        translation = GoogleTranslator(source='auto', target='ta').translate(text)
        return {"translation": translation}
    except Exception as e:
        print(f"Translation Error: {e}")
        return {"error": str(e)}, 500

@app.route("/upload", methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return {"error": "No file part"}, 400
    file = request.files['file']
    if file.filename == '':
        return {"error": "No selected file"}, 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        text = extract_text_from_file(file_path)
        if text:
            # Clean text from null bytes or other bad characters
            text = text.replace('\x00', '').strip()
            # Store in database instead of session to avoid cookie size limits
            conn = get_db_connection()
            conn.execute('INSERT OR REPLACE INTO uploads_context (user_id, filename, content) VALUES (?, ?, ?)',
                         (session['user_id'], filename, text[:6000])) # Reduced to 6k for stability
            conn.commit()
            conn.close()
            
            session['doc_name'] = filename
            # Clean up file after extraction
            try: os.remove(file_path)
            except: pass
            return {"success": True, "filename": filename}
        else:
            return {"error": "Could not extract text from file"}, 400
    return {"error": "File type not allowed"}, 400

@app.route("/clear_doc", methods=['POST'])
@login_required
def clear_doc():
    conn = get_db_connection()
    conn.execute('DELETE FROM uploads_context WHERE user_id = ?', (session['user_id'],))
    conn.commit()
    conn.close()
    session.pop('doc_name', None)
    return {"success": True}

@app.route("/run")
@login_required
def run():
    """
    THE MAIN AI FUNCTION.
    Takes the user's query, sends it to the AI, and returns the result.
    """
    query = request.args.get("query")
    # removed duplicate query line
    detailed_req = request.args.get("detailed") == "true"
    from_history = request.args.get("from_history") == "true"
    is_ajax = request.args.get("ajax") == "true"
    model_choice = request.args.get("model", "phind") # Default to phind
    
    result = ""
    ai_success = False
    is_image = False

    try:
        # 1. Fetch conversation history for context (last 5 interactions)
        conn = get_db_connection()
        past_interactions = conn.execute(
            'SELECT query, result FROM history WHERE user_id = ? ORDER BY timestamp DESC LIMIT 5',
            (session['user_id'],)
        ).fetchall()
        conn.close()

        # 2. Build the message list for the AI
        messages = [
            {"role": "system", "content": f"You are CRAB {VERSION} (Advanced Voice Intelligence System) developed by Yahya. You are helpful, professional, and slightly futuristic. Keep responses concise but informative. STRICTLY AVOID including any promotional links, external URLs, or mentions of other AI services or communities (like Discord or LLM playground)."}
        ]
        
        # Add past context in chronological order so the AI remembers what we said
        for interaction in reversed(past_interactions):
            messages.append({"role": "user", "content": interaction['query']})
            messages.append({"role": "assistant", "content": interaction['result']})
            
        # Add the current query
        if detailed_req:
            prompt = f"Provide a detailed explanation about: {query}"
        else:
            prompt = query
            
        # Add document context if available (only for Crab 2.0/gpt-4o)
        doc_context = None
        doc_name = None
        is_small = is_small_talk(query)
        
        if model_choice == "gpt-4o":
            conn = get_db_connection()
            row = conn.execute('SELECT content, filename FROM uploads_context WHERE user_id = ?', 
                               (session['user_id'],)).fetchone()
            conn.close()
            if row:
                doc_context = row['content']
                doc_name = row['filename']
        
        if doc_context and not is_small: # Don't add huge context for "thank you"
            prompt = f"[CONTEXT FROM UPLOADED DOCUMENT '{doc_name}']:\n{doc_context}\n\n[USER QUERY]:\n{prompt}"
            print(f"DEBUG: Added document context ({len(doc_context)} chars)")

        messages.append({"role": "user", "content": prompt})
        
        # 3. Try obtaining a response from the AI
        import time
        start_time = time.time()
        
        # Determine providers based on model_choice
        if model_choice == "gpt-4o":
            if is_small:
                # FAST PATH for small talk - prioritize faster providers
                providers = [
                    (g4f.Provider.PollinationsAI, "gemini"),
                    (g4f.Provider.BlackboxPro, g4f.models.gpt_4o),
                    (g4f.Provider.ApiAirforce, g4f.models.gemini_2_0_flash),
                    (g4f.Provider.AnyProvider, g4f.models.gemini),
                ]
            else:
                # CRAB 2.0: Gemini 2.0 Flash / GPT-4o Standard Path
                providers = [
                    (g4f.Provider.PollinationsAI, "gemini"),
                    (g4f.Provider.ApiAirforce, g4f.models.gemini_2_0_flash),
                    (g4f.Provider.BlackboxPro, g4f.models.gpt_4o),
                    (g4f.Provider.AnyProvider, g4f.models.gemini_2_0_flash),
                    (g4f.Provider.AnyProvider, g4f.models.gpt_4o),
                ]
        elif model_choice == "phind" or not model_choice:
            # CRAB 1.0: GPT-4 (via g4f as requested)
            providers = [
                (g4f.Provider.ApiAirforce, g4f.models.gpt_4),
                (g4f.Provider.BlackboxPro, g4f.models.gpt_4),
                (g4f.Provider.AnyProvider, g4f.models.gpt_4),
            ]
            if is_small: # Always use gemini for small talk speed
                providers.insert(0, (g4f.Provider.PollinationsAI, "gemini"))
        elif model_choice == "gpt-4o-mini":
            providers = [
                (g4f.Provider.PollinationsAI, g4f.models.gpt_4o_mini),
                (g4f.Provider.ApiAirforce, g4f.models.gpt_4o_mini),
            ]
        else: # Generic fallback
            providers = [
                (g4f.Provider.PollinationsAI, "gemini"),
                (g4f.Provider.AnyProvider, g4f.models.default),
                (g4f.Provider.BlackboxPro, g4f.models.gpt_4o),
            ]

        # 3. Try obtaining a response from Groq first (Primary/Reliable)
        try:
            print("DEBUG: Attempting Groq (Primary)...")
            # Select model for Groq
            groq_model = "llama-3.3-70b-versatile" # Default high-quality model
            if model_choice == "gpt-4o-mini":
                groq_model = "llama-3.1-8b-instant" # Fast/light model
            
            chat_completion = client.chat.completions.create(
                messages=messages,
                model=groq_model,
            )
            if chat_completion.choices[0].message.content:
                result = chat_completion.choices[0].message.content
                ai_success = True
                print("DEBUG: Groq success!")
        except Exception as groq_error:
            print(f"DEBUG: Groq failed or rate-limited: {groq_error}")

        # 4. Fallback to g4f providers if Groq failed
        if not ai_success:
            def get_ai_response(p_item):
                p_prov, p_mod = p_item
                try:
                    p_name = getattr(p_prov, '__name__', 'Unknown')
                    print(f"DEBUG: Parallel escape attempt with {p_name}...")
                    resp = g4f.ChatCompletion.create(
                        model=p_mod,
                        messages=messages,
                        provider=p_prov
                    )
                    if resp and len(resp.strip()) > 0:
                        return resp
                except: pass
                return None

            # Try first 3 providers in parallel for maximum speed
            with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                # Parallelize based on working providers from research
                test_providers = providers[:3]
                if g4f.Provider.OperaAria not in [p[0] for p in providers]:
                    test_providers.append((g4f.Provider.OperaAria, g4f.models.gpt_4o))
                
                future_to_prov = {executor.submit(get_ai_response, p): p for p in test_providers}
                for future in concurrent.futures.as_completed(future_to_prov):
                    res = future.result()
                    if res:
                        result = res
                        ai_success = True
                        elapsed = time.time() - start_time
                        print(f"DEBUG: Parallel fallback success in {elapsed:.2f}s!")
                        break

            # Sequential fallback for the rest
            if not ai_success:
                for provider, model in providers[3:]:
                    try:
                        p_name = getattr(provider, '__name__', 'Unknown')
                        print(f"DEBUG: Sequential fallback with {p_name}...")
                        response = g4f.ChatCompletion.create(
                            model=model,
                            messages=messages,
                            provider=provider
                        )
                        if response and len(response.strip()) > 0:
                            result = response
                            ai_success = True
                            break
                    except:
                        continue

        # Final Fallback for Crab 2.0 if everything else fails: Try Gemini search/fast
        if not ai_success and model_choice == "gpt-4o":
            try:
                print("DEBUG: Gemini 2.0 Flash failed, falling back to gemini-fast...")
                response = g4f.ChatCompletion.create(
                    model="gemini",
                    messages=messages,
                    provider=g4f.Provider.PollinationsAI
                )
                if response and len(response.strip()) > 0:
                    result = response
                    ai_success = True
            except: pass
        
        if result:
            # Filter out promotional lines and specific phrases from the AI response
            promotional_keywords = [
                "llmplayground.net",
                "Want best roleplay experience",
                "Join our Discord",
                "gpt4free",
                "Blackbox AI",
                "op.wtf",
                "proxies",
                "cheaper than the market",
                "discord.gg",
                "t.me/",
            ]
            
            lines = result.split('\n')
            filtered_lines = []
            for line in lines:
                if any(kw.lower() in line.lower() for kw in promotional_keywords):
                    continue
                filtered_lines.append(line)
            
            result = '\n'.join(filtered_lines).strip()

        if not ai_success:
            result = "I'm having trouble connecting to the AI service right now. Please try again in a moment."

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error in run: {e}")
        result = f"An unexpected error occurred: {str(e)}"
        ai_success = False

    # 4. Save the interaction to the database (History)
    history_id = None
    chat_id = request.args.get("chat_id")
    # Handle case where JS sends "null" string
    if chat_id == "null" or not chat_id: chat_id = None

    if ai_success and not from_history:
        conn = get_db_connection()
        
        # Create a new chat conversation if one doesn't exist
        if not chat_id:
            # Title is first 30 chars of query
            title = query[:30] + "..." if len(query) > 30 else query
            cursor = conn.execute('INSERT INTO chats (user_id, title) VALUES (?, ?)', (session['user_id'], title))
            chat_id = cursor.lastrowid
            conn.commit()

        cursor = conn.execute('INSERT INTO history (user_id, chat_id, query, result, doc_name) VALUES (?, ?, ?, ?, ?)', 
                             (session['user_id'], chat_id, query, result, session.get('doc_name')))
        history_id = cursor.lastrowid
        conn.commit()
        conn.close()

    # If this was a background request (AJAX), return JSON
    if is_ajax:
        return {
            "query": query,
            "result": result,
            "ai_success": ai_success,
            "history_id": history_id,
            "chat_id": chat_id,
            "show_export": is_export_request(query)
        }

    # Otherwise refresh the page
    return redirect(url_for('app_interface', chat_id=chat_id))

# -------------------------
# HISTORY ADJUSTMENT
# -------------------------

@app.route("/export/<int:msg_id>/<file_format>")
@login_required
def export_message(msg_id, file_format):
    """
    Export a specific AI message as PDF, DOCX, or TXT.
    """
    conn = get_db_connection()
    item = conn.execute('SELECT result, timestamp FROM history WHERE id = ? AND user_id = ?', 
                         (msg_id, session['user_id'])).fetchone()
    conn.close()

    if not item:
        return "Message not found", 404

    text = item['result']
    # Clean markdown if possible for cleaner documents
    clean_text = re.sub(r'(\*\*|__|`|#)', '', text)
    
    timestamp = item['timestamp'].replace(':', '-').replace(' ', '_')
    filename = f"CRAB_Response_{timestamp}"

    if file_format == 'txt':
        return send_file(
            io.BytesIO(text.encode('utf-8')),
            mimetype='text/plain',
            as_attachment=True,
            download_name=f"{filename}.txt"
        )

    elif file_format == 'docx':
        doc = docx.Document()
        doc.add_heading('CRAB AI Response', 0)
        doc.add_paragraph(f"Timestamp: {item['timestamp']}")
        doc.add_paragraph("-" * 20)
        doc.add_paragraph(text)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{filename}.docx"
        )

    elif file_format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        # fpdf2 supports unicode but needs a font that has those glyphs
        # For now, we'll try to use standard fonts and handle basic text
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(40, 10, "CRAB AI Response")
        pdf.ln(10)
        
        pdf.set_font("helvetica", "", 10)
        pdf.cell(40, 10, f"Timestamp: {item['timestamp']}")
        pdf.ln(10)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        
        pdf.set_font("helvetica", "", 11)
        # Use multi_cell for wrapping text
        # We need to sanitize text for latin-1 if using standard fonts
        safe_text = text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 7, safe_text)
        
        file_stream = io.BytesIO(pdf.output())
        
        return send_file(
            file_stream,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"{filename}.pdf"
        )

    return "Invalid format", 400

@app.route("/delete_history/<int:id>", methods=["POST"])
@login_required
def delete_history(id):
    # Deletes a single message
    conn = get_db_connection()
    conn.execute('DELETE FROM history WHERE id = ? AND user_id = ?', (id, session['user_id']))
    conn.commit()
    conn.close()
    return "OK"

@app.route("/delete_chat/<int:id>", methods=["POST"])
@login_required
def delete_chat(id):
    # Deletes an entire conversation
    conn = get_db_connection()
    conn.execute('DELETE FROM history WHERE chat_id = ? AND user_id = ?', (id, session['user_id']))
    conn.execute('DELETE FROM chats WHERE id = ? AND user_id = ?', (id, session['user_id']))
    conn.commit()
    conn.close()
    return "OK"

@app.route("/delete_all_history", methods=["POST"])
@login_required
def delete_all_history():
    # Deletes EVERYTHING (Danger Zone)
    conn = get_db_connection()
    conn.execute('DELETE FROM history WHERE user_id = ?', (session['user_id'],))
    conn.execute('DELETE FROM chats WHERE user_id = ?', (session['user_id'],))
    conn.commit()
    conn.close()
    return "OK"

# Initialize the DB when the app starts
init_db()

if __name__ == "__main__":
    # If frozen, debug must be False to avoid reloader issues
    if getattr(sys, 'frozen', False):
        # Open browser automatically
        def open_browser():
            time.sleep(1.5)
            webbrowser.open("http://127.0.0.1:5000")
        
        threading.Thread(target=open_browser).start()
        app.run(debug=False)
    else:
        app.run(debug=True)
